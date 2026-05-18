from docx import Document
from pathlib import Path


DOCS_DIR = Path("generated_docs")
DOCS_DIR.mkdir(exist_ok=True)


def create_docx(process_id, client_name, document_type, text):

    document = Document()

    document.add_heading("DocFlow AI", level=1)

    document.add_heading(document_type, level=2)

    document.add_paragraph(f"Cliente: {client_name}")

    document.add_paragraph("")

    document.add_paragraph(text)

    filename = f"documento_{process_id}.docx"

    filepath = DOCS_DIR / filename

    document.save(filepath)

    return str(filepath)